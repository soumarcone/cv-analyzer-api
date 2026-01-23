from io import BytesIO
from docx import Document
from app.core.config import settings


def extract_text_from_docx_bytes(data: bytes) -> tuple[str, dict]:
    """Extract text content from DOCX file bytes.
    
    Args:
        data: Raw bytes of the DOCX file.
    
    Returns:
        tuple: A tuple containing:
            - str: Extracted text from all paragraphs.
            - dict: Metadata with paragraph count.
    
    Raises:
        ValueError: If DOCX has too many paragraphs.
    """
    doc = Document(BytesIO(data))
    
    # Validate paragraph count
    para_count = len(doc.paragraphs)
    max_paras = settings.app.max_docx_paragraphs
    
    if para_count > max_paras:
        raise ValueError(
            f"DOCX has too many paragraphs: {para_count} (max allowed: {max_paras})"
        )
    
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    full_text = "\n".join(paragraphs).strip()
    meta = {"paragraphs": len(paragraphs)}
    return full_text, meta