"""Unit tests for DOCX extractor with paragraph limit validation."""

import io
import pytest
from docx import Document
from app.utils.docx_extractor import extract_text_from_docx_bytes
from app.core.config import settings


class TestDOCXExtractorParagraphLimits:
    """Test DOCX paragraph limit enforcement."""

    def _create_docx_with_paragraphs(self, num_paragraphs: int) -> bytes:
        """Create a DOCX with specified number of paragraphs.
        
        Args:
            num_paragraphs: Number of paragraphs to create.
            
        Returns:
            bytes: DOCX file content.
        """
        doc = Document()
        
        for i in range(num_paragraphs):
            doc.add_paragraph(f"This is paragraph number {i + 1}.")
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def test_extract_docx_within_limit(self):
        """DOCX with paragraphs within limit should extract successfully."""
        # Create DOCX with 50 paragraphs (well below default limit of 500)
        docx_data = self._create_docx_with_paragraphs(50)
        
        text, meta = extract_text_from_docx_bytes(docx_data)
        
        assert isinstance(text, str)
        assert meta["paragraphs"] == 50
        assert "paragraph number 1" in text.lower()

    def test_extract_docx_at_exact_limit(self):
        """DOCX with exactly max paragraphs should extract successfully."""
        max_paras = settings.app.max_docx_paragraphs
        docx_data = self._create_docx_with_paragraphs(max_paras)
        
        text, meta = extract_text_from_docx_bytes(docx_data)
        
        assert isinstance(text, str)
        assert meta["paragraphs"] == max_paras

    def test_extract_docx_exceeds_limit(self):
        """DOCX exceeding paragraph limit should raise ValueError."""
        max_paras = settings.app.max_docx_paragraphs
        docx_data = self._create_docx_with_paragraphs(max_paras + 1)
        
        with pytest.raises(ValueError) as exc_info:
            extract_text_from_docx_bytes(docx_data)
        
        error_msg = str(exc_info.value)
        assert "too many paragraphs" in error_msg.lower()
        assert str(max_paras + 1) in error_msg
        assert str(max_paras) in error_msg

    def test_extract_docx_far_exceeds_limit(self):
        """DOCX with many paragraphs over limit should be rejected."""
        max_paras = settings.app.max_docx_paragraphs
        # Create DOCX with double the limit
        docx_data = self._create_docx_with_paragraphs(max_paras * 2)
        
        with pytest.raises(ValueError) as exc_info:
            extract_text_from_docx_bytes(docx_data)
        
        error_msg = str(exc_info.value)
        assert "too many paragraphs" in error_msg.lower()

    def test_extract_docx_empty_document(self):
        """Empty DOCX (0 paragraphs) should extract successfully."""
        # Document() creates a doc with 1 empty paragraph by default
        doc = Document()
        # Clear all paragraphs
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_data = buffer.read()
        
        text, meta = extract_text_from_docx_bytes(docx_data)
        
        # Should succeed with 0 paragraphs
        assert meta["paragraphs"] == 0
        assert text == ""
